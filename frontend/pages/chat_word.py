import streamlit as st
import pandas as pd
from docx import Document
from groq import Groq
from dotenv import load_dotenv
import os
from components.navbar import show_navbar

load_dotenv()  # This loads the .env file into environment variables

st.set_page_config(page_title="DataMate • Chat Document", page_icon="https://img.icons8.com/?size=100&id=112370&format=png&color=000000")
show_navbar()

# Title 
st.markdown(
    """
    <div style='display: flex; align-items: center; gap: 20px; margin-bottom: 1rem;'>
        <img src='https://img.icons8.com/?size=100&id=11571&format=png&color=000000' width='50' height='50' style='margin-left: 10px;'/>
        <h1 style='margin: 0;'>Chat with Word Document</h1>
    </div>
    """,
    unsafe_allow_html=True
)

# Initialize Groq
try:
    groq_api_key = st.secrets["GROQ_API_KEY"]
    client = Groq(api_key=groq_api_key)
except KeyError:
    st.error("❌ GROQ_API_KEY is missing from Streamlit secrets. Please add it to your `.streamlit/secrets.toml`.")
    st.stop()

# Initialize chat state
if "word_messages" not in st.session_state:
    st.session_state["word_messages"] = [{"role": "assistant", "content": "Ask your questions!"}]
if "word_chat_history" not in st.session_state:
    st.session_state["word_chat_history"] = []

# File uploader
uploaded_word = st.file_uploader("Upload a Word document", type=["docx"])
if uploaded_word:
    text_chunks = []
    doc = Document(uploaded_word)
    total_paragraphs = len(doc.paragraphs)
    progress = st.progress(0, text="Starting to read the Word document...")

    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            text_chunks.append(para.text.strip())
        progress.progress((idx + 1) / total_paragraphs, text=f"Loaded paragraph {idx + 1} of {total_paragraphs}")

    # Combine content
    word_text = "\n\n".join(text_chunks[:])  # You may limit paragraphs for token efficiency

    # Show conversation history
    for msg in st.session_state["word_messages"]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    user_query = st.chat_input("Ask a question about this Word document...")

    if user_query:
        st.session_state["word_messages"].append({"role": "user", "content": user_query})
        with st.chat_message("user"):
            st.markdown(user_query)

        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                prompt = f"""
You are a helpful assistant. A user uploaded a Word document. Based on the following extracted content, answer their question in plain English.

Word document content:
{word_text}

User's question: {user_query}

Answer:"""
                try:
                    response = client.chat.completions.create(
                        model="llama3-8b-8192",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.4
                    )
                    answer = response.choices[0].message.content.strip()
                    st.markdown("**Answer:**")
                    st.write(answer)
                    st.session_state["word_messages"].append({"role": "assistant", "content": answer})
                    st.session_state["word_chat_history"].append({"Q": user_query, "A": answer})
                except Exception as e:
                    error_str = str(e)
                    if "Request too large" in error_str or "'code': 'rate_limit_exceeded'" in error_str:
                        error_msg = "❌ File too long. The document exceeds the model's processing limit. Please upload a shorter file or split it into smaller parts."
                    else:
                        error_msg = f"❌ Error: {e}"
                    st.error(error_msg)
                    st.session_state["word_messages"].append({"role": "assistant", "content": error_msg})
                    
    # Clear chat history
    if st.sidebar.button("🗑️ Clear Chat History"):
        st.session_state["word_messages"] = [{"role": "assistant", "content": "Ask your questions!"}]
        st.session_state["word_chat_history"] = []
        st.rerun()

    # Chat download option
    if st.session_state["word_chat_history"]:
        if st.sidebar.download_button("📥 Download Chat",
            data=pd.DataFrame(st.session_state["word_chat_history"]).to_csv(index=False),
            file_name="word_chat_history.csv",
            mime="text/csv"):
            st.sidebar.success("Downloaded!")
else:
    st.info("Please upload a Word document to begin.")
